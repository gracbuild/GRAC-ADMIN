from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).with_name("GRAC_Control_Management_User_Manual.docx")
NAVY = "0B2545"; BLUE = "2E74B5"; DARK_BLUE = "1F4D78"; MUTED = "667085"
HEADER_FILL = "E8EEF5"; LIGHT_FILL = "F4F6F9"; GOLD = "7A5A00"

SCREENS = [
    ("Authorities", "Maintain issuing and supervisory bodies.", "name, description, jurisdiction, website, status"),
    ("Regulatory Artifacts", "Maintain regulations, standards, laws, directives, circulars, guidelines, and accreditation programs.", "authorityId, name, description, category, industry, jurisdiction, status"),
    ("Releases", "Maintain published versions of regulatory artifacts.", "artifactId, version, effectiveDate, endDate, releaseNotes, status"),
    ("Framework Source Structure", "Preserve native hierarchy and terminology for each release.", "releaseId, parentNodeId, nodeLevel, nodeType, reference, title, description, displayOrder"),
    ("Controls", "Maintain normalized reusable compliance objectives.", "code, name, description, objective"),
    ("Requirements", "Maintain atomic, independently assessable compliance expectations.", "code, name, statement, objective, status"),
    ("Obligations", "Maintain release-specific execution expectations.", "requirementId, releaseId, mandatory, frequencyType, frequencyValue, frequencyUnit, triggerCondition, dueWithin, evidenceRequired, evidenceType, retentionRequirement, severity"),
    ("Control-Requirement Mapping", "Associate reusable requirements with controls.", "controlId, requirementId"),
    ("Source Structure-Control Mapping", "Link native framework leaf references to normalized controls; release association is derived from the selected source node.", "structureNodeId, controlId"),
    ("Applicability Rules", "Maintain metadata-driven organizational scope rules.", "artifactId, releaseId, name, expression, priority, outcome"),
    ("Change Management", "Record regulatory repository lifecycle events.", "entityType, entityId, changeType, summary, effectiveDate, severity"),
    ("Impact Analysis", "Record repository and organizational impacts.", "changeEventId, impactedEntityType, impactedEntityId, organizationId, summary, recommendedAction"),
    ("Notification View", "Create and review organization notifications.", "impactAnalysisId, organizationId, type, subject, message, severity, recommendedAction"),
]

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None: tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths): col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")
            set_cell_margin(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def mark_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name; run._element.rPr.rFonts.set(qn("w:ascii"), name); run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic

def style_para(p, before=0, after=6, line=1.25):
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = line

def add_para(doc, text="", bold=False, color=None, size=11, align=None, before=0, after=6, italic=False):
    p = doc.add_paragraph(); style_para(p, before, after)
    if align is not None: p.alignment = align
    r = p.add_run(text); set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); style_para(p, after=4)
    set_font(p.add_run(text), size=11)

def add_step(doc, text):
    p = doc.add_paragraph(style="List Number"); style_para(p, after=4)
    set_font(p.add_run(text), size=11)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    return p.add_run(text)

def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"; set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], HEADER_FILL)
        p = table.rows[0].cells[i].paragraphs[0]; style_para(p, after=0, line=1.0)
        set_font(p.add_run(h), size=9.5, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]; style_para(p, after=0, line=1.0)
            set_font(p.add_run(str(value)), size=9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table

def add_callout(doc, title, text, fill=LIGHT_FILL):
    t = doc.add_table(rows=1, cols=1); set_table_geometry(t, [9360]); mark_header_row(t.rows[0]); set_cell_shading(t.cell(0,0), fill)
    p=t.cell(0,0).paragraphs[0]; style_para(p, after=2)
    set_font(p.add_run(title + ": "), size=10.5, bold=True, color=DARK_BLUE)
    set_font(p.add_run(text), size=10.5)
    doc.add_paragraph()

doc = Document()
sec = doc.sections[0]
sec.top_margin=Inches(1); sec.bottom_margin=Inches(1); sec.left_margin=Inches(1); sec.right_margin=Inches(1)
sec.header_distance=Inches(.492); sec.footer_distance=Inches(.492)

styles=doc.styles
normal=styles["Normal"]; normal.font.name="Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [("Heading 1",16,BLUE,18,10),("Heading 2",13,BLUE,14,7),("Heading 3",12,DARK_BLUE,10,5)]:
    s=styles[name]; s.font.name="Calibri"; s._element.rPr.rFonts.set(qn("w:ascii"),"Calibri"); s._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri"); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
for name in ["List Bullet","List Number"]:
    styles[name].font.name="Calibri"; styles[name].font.size=Pt(11); styles[name].paragraph_format.space_after=Pt(4); styles[name].paragraph_format.line_spacing=1.25

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; style_para(header,after=0,line=1)
set_font(header.add_run("GRAC | Repository Management User Manual"), size=9, color=MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; style_para(footer,after=0,line=1)
set_font(footer.add_run("Standalone review module | User Manual"), size=9, color=MUTED)

# Cover
add_para(doc,"GRAC",bold=True,color=BLUE,size=14,after=48)
add_para(doc,"REPOSITORY MANAGEMENT",bold=True,color=NAVY,size=28,after=4)
add_para(doc,"User Manual",color=DARK_BLUE,size=18,after=10)
add_para(doc,"Regulatory Intelligence Repository & Change Management Engine",color=MUTED,size=12,after=36)
add_table(doc,["Document","Details"],[
    ("Version","1.0"),("Module","Standalone ControlManagement review build"),("Audience","Repository administrators, compliance analysts, reviewers"),("Prepared","June 2026")
],[2500,6860])
add_callout(doc,"Purpose","This manual explains how to navigate and use the standalone GRAC Repository Management module created for functional review.")
add_para(doc,"The module is designed as a configurable repository. Framework-specific names, taxonomies, releases, obligations, and applicability rules are maintained as data.",italic=True,color=MUTED,size=10.5)
doc.add_page_break()

add_heading(doc,"Contents",1)
for item in ["1. Module Overview","2. Access and Navigation","3. Common Screen Workflow","4. Recommended Onboarding Sequence","5. Screen-by-Screen Guide","6. Audit and Traceability","7. Status and Change Guidance","8. Review-Build Notes","9. Quick Reference"]:
    add_para(doc,item,after=3)
doc.add_page_break()

add_heading(doc,"1. Module Overview",1)
add_para(doc,"Repository Management is the GRAC repository for regulatory intelligence. It maintains authorities, regulatory artifacts, releases, native source structures, normalized controls, atomic requirements, release-specific obligations, applicability rules, change events, impact records, and notifications.")
add_heading(doc,"Repository Relationship",2)
for step in ["Authority","Regulatory Artifact","Release","Framework Source Structure","Source Structure-Control Mapping","Control","Control-Requirement Mapping","Requirement","Regulatory Obligation"]:
    add_para(doc,step,bold=True,color=DARK_BLUE,after=2)
add_callout(doc,"Design principle","Controls and requirements are normalized and reusable. Obligations remain release-specific because execution expectations can differ across publishers and versions.")
add_heading(doc,"What the module avoids",2)
add_bullet(doc,"No named regulatory framework is hardcoded into business logic.")
add_bullet(doc,"No physical delete is required for repository records.")
add_bullet(doc,"No native source terminology is discarded during normalization.")

add_heading(doc,"2. Access and Navigation",1)
add_para(doc,"Open the Repository Management web application. The left sidebar provides access to the overview dashboard, all repository management screens, and the audit traceability view.")
add_table(doc,["Dashboard Item","Meaning"],[
    ("Repository areas","Number of configurable management screens available."),
    ("Traceability","Indicates that soft delete and immutable audit history are part of the design."),
    ("Framework logic","Confirms that onboarding behavior is data-driven."),
],[2200,7160])
add_heading(doc,"Sidebar Navigation",2)
add_para(doc,"Select any sidebar item to open its repository workspace. Use Overview to return to the landing dashboard.")

add_heading(doc,"3. Common Screen Workflow",1)
add_para(doc,"Each management screen follows the same interaction pattern so users can move between repository areas without learning a new layout.")
add_table(doc,["Control","How to use it"],[
    ("Search repository","Type a keyword to filter the current list."),
    ("All statuses","Filter records by Active, Draft, or Retired status."),
    ("Refresh","Reload the current repository list from the API."),
    ("Add record","Open the record entry dialog."),
    ("View","Review the selected row. Detailed view behavior is part of the next integration phase."),
],[2200,7160])
add_heading(doc,"Adding a Record",2)
for text in ["Select the required repository area from the sidebar.","Choose Add record.","Enter your user identifier in Entered by.","Enter the record metadata as valid JSON.","Choose Save.","Review the success message and refresh the list if required."]:
    add_step(doc,text)
add_callout(doc,"Form controls","Each repository area uses a dedicated Add/Edit form with ordinary text fields, text areas, dates, checkboxes, and API-fed dropdowns. Mapping screens use multi-select controls where several relationships can be created together.",fill="FFF8E8")
add_heading(doc,"Example: Add an Authority",2)
add_para(doc,"Enter the authority name, description, jurisdiction, website, and status. The module generates the internal authority code automatically.",size=10.5,color=DARK_BLUE)

add_heading(doc,"4. Recommended Onboarding Sequence",1)
add_para(doc,"Use the following sequence when onboarding a new regulatory source. This preserves referential integrity and makes downstream mappings easier to review.")
for text in ["Create the Authority.","Create the Regulatory Artifact under the authority.","Create one or more Releases for the artifact.","Build the Framework Source Structure for each release.","Create or reuse normalized Controls.","Create or reuse atomic Requirements.","Map releases to controls.","Map native source structure nodes to controls.","Map controls to requirements.","Create release-specific Obligations.","Configure Applicability Rules.","Record Changes, Impact Analysis, and Notifications when a regulatory change is identified."]:
    add_step(doc,text)
add_callout(doc,"Important","Reuse existing controls and requirements whenever they express the same normalized compliance expectation. Do not create framework-specific duplicates unless the expectation is materially different.")

add_heading(doc,"5. Screen-by-Screen Guide",1)
for title, purpose, fields in SCREENS:
    add_heading(doc,title,2)
    add_para(doc,purpose,after=4)
    add_table(doc,["Entry metadata","Guidance"],[
    ("Form fields",fields),
        ("Good practice","Use stable codes and meaningful descriptions. Retire superseded records instead of deleting them."),
    ],[1900,7460])

add_heading(doc,"6. Audit and Traceability",1)
add_para(doc,"Use Audit Traceability to review the immutable repository history. Audit records capture the entity type, entity identifier, action type, status, user, and timestamp.")
add_table(doc,["Audit Column","Description"],[
    ("EntityType","Repository area affected by the action."),
    ("EntityId","Identifier of the affected record."),
    ("ActionType","Action recorded by the stored procedure, such as SAVE or RETIRE."),
    ("Status","Audit row status. Audit records remain immutable."),
    ("EnteredBy","User identifier supplied with the transaction."),
    ("EnteredDt","UTC timestamp recorded by the database."),
],[1900,7460])
add_callout(doc,"Traceability rule","Audit trace rows are append-only. The database trigger rejects updates and deletes against the audit table.")

add_heading(doc,"7. Status and Change Guidance",1)
add_table(doc,["Area","Typical states","Guidance"],[
    ("Authorities","Active, Inactive","Mark inactive when an authority should no longer be used for onboarding."),
    ("Artifacts","Active, Retired","Retire artifacts without removing historical releases."),
    ("Releases","Draft, Active, Retired, Archived","Multiple releases may remain active when required."),
    ("Controls","Active, Retired","Retire only when the normalized objective is no longer applicable."),
    ("Requirements","Active, Retired","Track lifecycle events through Change Management and Audit Traceability."),
    ("Notifications","Pending, Archived","Archive notifications after operational handling."),
],[1450,2100,5810])

add_heading(doc,"8. Review-Build Notes",1)
add_para(doc,"This manual describes the standalone first-review module. The following items remain explicit decisions for integration into the main GRAC platform:")
for text in ["Authentication and role-based authorization.","Approval workflow for repository publishing.","Organization directory integration and client applicability execution.","Notification delivery channels and templates."]:
    add_bullet(doc,text)
add_callout(doc,"Current behavior","All requested repository screens are present. Reads and stored-procedure-backed writes are implemented for the requested areas. The list UI can be reviewed independently after the API connection string is configured.")

add_heading(doc,"9. Quick Reference",1)
add_table(doc,["Need to...","Open this screen"],[
    ("Register a publisher","Authorities"),("Add a regulation, standard, law, or program","Regulatory Artifacts"),
    ("Create a version","Releases"),("Preserve native references","Framework Source Structure"),
    ("Define a normalized objective","Controls"),("Define an atomic expectation","Requirements"),
    ("Record execution expectations","Obligations"),("Link controls and requirements","Control-Requirement Mapping"),
    ("Derive release-control coverage","Source Structure-Control Mapping"),("Link native references and controls","Source Structure-Control Mapping"),
    ("Determine scope through metadata","Applicability Rules"),("Record a regulatory change","Change Management"),
    ("Assess effect on repository or organizations","Impact Analysis"),("Prepare organization communication","Notification View"),
    ("Review historical actions","Audit Traceability"),
],[3300,6060])
add_para(doc,"End of manual",bold=True,color=MUTED,size=10,align=WD_ALIGN_PARAGRAPH.CENTER,before=18)

doc.save(OUT)
print(OUT)
